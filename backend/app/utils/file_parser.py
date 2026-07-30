import os
from pypdf import PdfReader
from docx import Document

def extract_text_from_file(file_path: str) -> str:
    """
    Extracts text from TXT, PDF, and DOCX files.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    _, ext = os.path.splitext(file_path.lower())
    
    if ext == ".txt":
        return extract_from_txt(file_path)
    elif ext == ".pdf":
        return extract_from_pdf(file_path)
    elif ext == ".docx":
        return extract_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {ext}. Only PDF, DOCX, and TXT are supported.")

def extract_from_txt(file_path: str) -> str:
    # Try different encodings
    encodings = ["utf-8", "latin-1", "utf-16"]
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode TXT file with standard encodings.")

def extract_from_pdf(file_path: str) -> str:
    try:
        reader = PdfReader(file_path)
        text_parts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        extracted_text = "\n".join(text_parts).strip()
        if not extracted_text:
            return "[Empty PDF file or scanned image without OCR text]"
        return extracted_text
    except Exception as e:
        raise ValueError(f"Failed to parse PDF file: {str(e)}")

def extract_from_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs]
        # Include tables text if any
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(" | ".join(row_text))
        
        full_text = "\n".join(paragraphs + table_text).strip()
        if not full_text:
            return "[Empty DOCX document]"
        return full_text
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX file: {str(e)}")


if __name__ == "__main__":
    # Small test
    import tempfile
    
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tf:
        tf.write(b"Hello world from Edunova AI test file!")
        temp_name = tf.name
        
    try:
        txt = extract_text_from_file(temp_name)
        print("TXT Parsing Test:", txt)
    finally:
        os.remove(temp_name)
